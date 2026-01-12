"""
File converters for extracting text from different file formats.
Consolidates: pdf-to-txt, word-to-txt, img-to-txt, rtf-to-txt, txt-passthrough

Uses multi-library fallback chains for maximum reliability:
- PDF: pymupdf (fitz) → pdfplumber → PyPDF2
- DOCX: python-docx → mammoth → XML extraction → docx2txt
- DOC: LibreOffice → antiword
"""

import asyncio
import logging
import os
import shutil
import subprocess
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, Tuple
from xml.etree import ElementTree as ET

import cv2
import numpy as np
import pytesseract
from striprtf.striprtf import rtf_to_text

from config import ServiceConfig, SupportedExtensions

logger = logging.getLogger("resume-extractor.converters")

# Configure tesseract path
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# Thread pool for CPU-bound operations
_thread_pool = ThreadPoolExecutor(max_workers=ServiceConfig.FILE_PROCESSING_CONCURRENCY)


class TextConverter:
    """Base class for text converters."""

    @staticmethod
    async def convert(file_path: str) -> str:
        """Convert file to text. Override in subclasses."""
        raise NotImplementedError


class PDFConverter(TextConverter):
    """
    Convert PDF files to text using multi-library fallback chain.

    Priority order:
    1. pymupdf (fitz) - Fastest, has OCR support for scanned PDFs
    2. pdfplumber - Best for tables and complex layouts
    3. PyPDF2 - Fallback for edge cases
    """

    @staticmethod
    def _extract_with_pymupdf(file_path: str) -> Tuple[str, str]:
        """
        Extract text using PyMuPDF (fastest, best for general use).
        Returns tuple of (extracted_text, method_used).
        """
        try:
            import fitz  # pymupdf

            text_content = []
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Try standard text extraction first
                text = page.get_text("text")

                # If page has minimal text, might be scanned - try OCR
                if len(text.strip()) < 50:
                    logger.debug(
                        f"Page {page_num} has minimal text ({len(text.strip())} chars), attempting OCR"
                    )
                    try:
                        # Get OCR text using pymupdf's built-in OCR
                        tp = page.get_textpage_ocr(language="eng", dpi=300)
                        text = page.get_text(textpage=tp)
                    except Exception as ocr_error:
                        logger.debug(f"OCR failed on page {page_num}: {ocr_error}")

                if text:
                    text_content.append(text)

            doc.close()

            final_text = "\n".join(text_content).strip()

            if final_text:
                return final_text, "pymupdf"
            else:
                raise ValueError("No text extracted with pymupdf")

        except Exception as e:
            logger.debug(f"pymupdf extraction failed for {file_path}: {e}")
            raise

    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> Tuple[str, str]:
        """
        Extract text using pdfplumber (best for tables and structured layouts).
        Returns tuple of (extracted_text, method_used).
        """
        try:
            import pdfplumber

            text_content = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Try table extraction first (resumes often have skill tables)
                    tables = page.extract_tables()

                    if tables:
                        for table in tables:
                            table_text = "\n".join(
                                [
                                    " | ".join(
                                        filter(None, [str(cell) if cell else "" for cell in row])
                                    )
                                    for row in table
                                    if row
                                ]
                            )
                            if table_text.strip():
                                text_content.append(table_text)

                    # Extract regular text with layout preservation
                    text = page.extract_text(x_tolerance=3, y_tolerance=3, layout=True)

                    if text:
                        text_content.append(text)

            final_text = "\n\n".join(text_content).strip()

            if final_text:
                return final_text, "pdfplumber"
            else:
                raise ValueError("No text extracted with pdfplumber")

        except Exception as e:
            logger.debug(f"pdfplumber extraction failed for {file_path}: {e}")
            raise

    @staticmethod
    def _extract_with_pypdf2(file_path: str) -> Tuple[str, str]:
        """
        Extract text using PyPDF2 (fallback for edge cases).
        Returns tuple of (extracted_text, method_used).
        """
        try:
            import PyPDF2

            text_content = []

            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)

                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

            final_text = "\n".join(text_content).strip()

            if final_text:
                return final_text, "pypdf2"
            else:
                raise ValueError("No text extracted with PyPDF2")

        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed for {file_path}: {e}")
            raise

    @staticmethod
    def _extract_text(file_path: str) -> str:
        """Extract text from PDF using fallback chain (blocking operation)."""
        methods = [
            ("pymupdf", PDFConverter._extract_with_pymupdf),
            ("pdfplumber", PDFConverter._extract_with_pdfplumber),
            ("pypdf2", PDFConverter._extract_with_pypdf2),
        ]

        last_error = None

        for method_name, method_func in methods:
            try:
                text, actual_method = method_func(file_path)

                # Validate extracted text quality - if too short, try next method
                if len(text.strip()) < 20:
                    logger.debug(
                        f"{method_name} extracted minimal text ({len(text)} chars) for {file_path}, trying next method"
                    )
                    continue

                logger.info(
                    f"PDF extraction success: {os.path.basename(file_path)} using {actual_method} ({len(text)} chars)"
                )
                return text

            except Exception as e:
                last_error = e
                continue

        # All methods failed
        logger.error(f"All PDF extraction methods failed for {file_path}. Last error: {last_error}")
        return ""

    @staticmethod
    async def convert(file_path: str) -> str:
        """Convert PDF to text asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(_thread_pool, PDFConverter._extract_text, file_path)


class WordConverter(TextConverter):
    """
    Convert Word documents (.doc, .docx) to text.

    For .docx files - fallback chain:
    1. python-docx - Best table support
    2. mammoth - Best for LLM processing (markdown output)
    3. XML extraction - Works on corrupted files
    4. docx2txt - Simple fallback

    For .doc files:
    1. LibreOffice conversion to .docx → then .docx chain
    2. antiword - Direct text extraction fallback
    """

    # Semaphore to limit concurrent .doc to .docx conversions
    _conversion_semaphore = asyncio.Semaphore(ServiceConfig.DOC_CONVERSION_CONCURRENCY)

    @staticmethod
    def _extract_with_python_docx(file_path: str) -> str:
        """
        Primary method: Extract text using python-docx (best for tables).
        """
        try:
            from docx import Document

            document = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables (critical for resumes with skill matrices)
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            result = "\n".join(text_parts)
            if result.strip():
                return result
            raise ValueError("No text extracted with python-docx")

        except Exception as e:
            logger.debug(f"python-docx extraction failed: {e}")
            raise

    @staticmethod
    def _extract_with_mammoth(file_path: str) -> str:
        """
        Secondary method: Convert to markdown using mammoth.
        Excellent for LLM processing - preserves semantic structure.
        """
        try:
            import mammoth

            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                text = result.value.strip()
                if text:
                    return text
                raise ValueError("No text extracted with mammoth")

        except Exception as e:
            logger.debug(f"mammoth extraction failed: {e}")
            raise

    @staticmethod
    def _extract_from_corrupted_docx(file_path: str) -> str:
        """
        Tertiary method: Direct XML extraction for corrupted files.
        """
        try:
            with zipfile.ZipFile(file_path) as z:
                with z.open("word/document.xml") as f:
                    xml_content = f.read()

            namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            tree = ET.fromstring(xml_content)
            paragraphs = tree.findall(".//w:p", namespaces)

            text_lines = []
            for para in paragraphs:
                texts = [node.text for node in para.findall(".//w:t", namespaces) if node.text]
                text_lines.append("".join(texts))

            result = "\n".join(filter(None, text_lines))
            if result.strip():
                return result
            raise ValueError("No text extracted from XML")

        except Exception as e:
            logger.debug(f"XML extraction failed for {file_path}: {e}")
            raise

    @staticmethod
    def _extract_with_docx2txt(file_path: str) -> str:
        """
        Fallback method: Simple text extraction with docx2txt.
        Fast but loses table structure.
        """
        try:
            import docx2txt

            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
            raise ValueError("No text extracted with docx2txt")

        except Exception as e:
            logger.debug(f"docx2txt extraction failed: {e}")
            raise

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """
        Extract text from .docx with intelligent fallback chain (blocking operation).
        """
        strategies = [
            ("python-docx", WordConverter._extract_with_python_docx),
            ("mammoth", WordConverter._extract_with_mammoth),
            ("xml-extraction", WordConverter._extract_from_corrupted_docx),
            ("docx2txt", WordConverter._extract_with_docx2txt),
        ]

        for strategy_name, strategy_func in strategies:
            try:
                text = strategy_func(file_path)
                if text and len(text.strip()) > 0:
                    logger.info(
                        f"DOCX extraction success: {os.path.basename(file_path)} using {strategy_name} ({len(text)} chars)"
                    )
                    return text
            except Exception:
                continue

        logger.error(f"All DOCX extraction strategies failed for {file_path}")
        return ""

    @staticmethod
    async def _convert_doc_to_docx(doc_path: str) -> Optional[str]:
        """
        Convert .doc to .docx using LibreOffice with best practices.
        Includes timeout and proper isolation for parallel processing.
        """
        output_dir = os.path.dirname(doc_path)
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(output_dir, base_name + ".docx")
        user_profile = f"/tmp/lo_profile_{uuid.uuid4().hex}"

        env = os.environ.copy()
        env["HOME"] = "/tmp"

        try:
            process = await asyncio.create_subprocess_exec(
                "soffice",
                "--headless",
                "--nofirststartwizard",
                "--norestore",  # Prevent session restoration (fixes race conditions)
                "--nologo",  # Skip splash screen (faster startup)
                f"-env:UserInstallation=file://{user_profile}",
                "--convert-to",
                "docx",
                "--outdir",
                output_dir,
                doc_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )

            # Add timeout to prevent hangs on corrupted files
            try:
                stdout, stderr = await asyncio.wait_for(
                    process.communicate(),
                    timeout=30.0,  # 30 second timeout
                )
            except asyncio.TimeoutError:
                process.kill()
                await process.wait()
                logger.error(f"LibreOffice conversion timeout for {doc_path}")
                return None

            # Clean up profile immediately
            shutil.rmtree(user_profile, ignore_errors=True)

            if os.path.exists(docx_path):
                logger.debug(f"Successfully converted {doc_path} to {docx_path}")
                return docx_path
            else:
                logger.warning(
                    f"LibreOffice conversion failed (rc={process.returncode}): {stderr.decode()}"
                )
                return None

        except Exception as e:
            logger.error(f"Error converting .doc to .docx: {e}")
            shutil.rmtree(user_profile, ignore_errors=True)
            return None

    @staticmethod
    def _extract_doc_with_antiword(doc_path: str) -> str:
        """
        Fallback: Extract text from .doc using antiword.
        Faster than LibreOffice, but limited formatting support.
        """
        try:
            result = subprocess.run(
                ["antiword", "-m", "UTF-8.txt", doc_path],
                capture_output=True,
                text=True,
                timeout=10.0,
            )

            if result.returncode == 0 and result.stdout.strip():
                logger.info(f"DOC extraction success: {os.path.basename(doc_path)} using antiword")
                return result.stdout.strip()
            else:
                raise RuntimeError(f"antiword returned {result.returncode}: {result.stderr}")

        except FileNotFoundError:
            logger.debug("antiword not installed, skipping fallback")
            raise
        except subprocess.TimeoutExpired:
            logger.warning(f"antiword timeout for {doc_path}")
            raise

    @staticmethod
    async def convert(file_path: str) -> str:
        """Convert Word document to text."""
        loop = asyncio.get_event_loop()

        # Handle .doc files (legacy binary format)
        if file_path.lower().endswith(".doc"):
            async with WordConverter._conversion_semaphore:
                # Try LibreOffice conversion first
                docx_path = await WordConverter._convert_doc_to_docx(file_path)
                if docx_path and os.path.exists(docx_path):
                    try:
                        text = await loop.run_in_executor(
                            _thread_pool, WordConverter._extract_from_docx, docx_path
                        )
                        return text
                    finally:
                        # Clean up converted file
                        try:
                            os.remove(docx_path)
                        except:
                            pass

                # Fallback to antiword if LibreOffice failed
                try:
                    text = await loop.run_in_executor(
                        _thread_pool, WordConverter._extract_doc_with_antiword, file_path
                    )
                    if text:
                        return text
                except Exception as e:
                    logger.debug(f"antiword fallback failed: {e}")

                logger.error(f"All DOC extraction strategies failed for {file_path}")
                return ""
        else:
            # Handle .docx files
            return await loop.run_in_executor(
                _thread_pool, WordConverter._extract_from_docx, file_path
            )


class ImageConverter(TextConverter):
    """Convert images to text using OCR (pytesseract + OpenCV)."""

    @staticmethod
    def _deskew_image(image: np.ndarray) -> np.ndarray:
        """Deskew an image using Hough Transform to correct text alignment."""
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            gray = cv2.bitwise_not(gray)
            coords = np.column_stack(np.where(gray > 0))

            if len(coords) == 0:
                return image

            angle = cv2.minAreaRect(coords)[-1]
            angle = -(90 + angle) if angle < -45 else -angle

            (h, w) = image.shape[:2]
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(
                image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
            )
            return rotated
        except Exception as e:
            logger.warning(f"Deskew failed: {e}, returning original image")
            return image

    @staticmethod
    def _preprocess_image(image: np.ndarray) -> np.ndarray:
        """Preprocess the image to enhance it for OCR."""
        try:
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

            # Resize the image for better OCR results
            scale_percent = 150
            width = int(gray.shape[1] * scale_percent / 100)
            height = int(gray.shape[0] * scale_percent / 100)
            resized = cv2.resize(gray, (width, height), interpolation=cv2.INTER_LINEAR)

            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(resized, (5, 5), 0)

            # Adaptive thresholding for binarization
            thresholded = cv2.adaptiveThreshold(
                blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )

            # Morphological operations to clean up noise
            kernel = np.ones((2, 2), np.uint8)
            processed = cv2.morphologyEx(thresholded, cv2.MORPH_OPEN, kernel)

            return processed
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image

    @staticmethod
    def _extract_text(file_path: str) -> str:
        """Extract text from image using OCR (blocking operation)."""
        try:
            image = cv2.imread(file_path)
            if image is None:
                logger.error(f"Failed to load image: {file_path}")
                return ""

            # Deskew and preprocess
            deskewed = ImageConverter._deskew_image(image)
            processed = ImageConverter._preprocess_image(deskewed)

            # OCR with optimized config
            custom_config = r"--psm 6 --oem 3"
            extracted_text = pytesseract.image_to_string(processed, config=custom_config)

            result = extracted_text.strip()
            if result:
                logger.info(
                    f"Image OCR success: {os.path.basename(file_path)} ({len(result)} chars)"
                )
            return result
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {e}")
            return ""

    @staticmethod
    async def convert(file_path: str) -> str:
        """Convert image to text using OCR asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(_thread_pool, ImageConverter._extract_text, file_path)


class RTFConverter(TextConverter):
    """Convert RTF files to text using striprtf."""

    @staticmethod
    def _extract_text(file_path: str) -> str:
        """Extract text from RTF file (blocking operation)."""
        try:
            # Try reading with different encodings
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
            rtf_content = None

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as rtf_file:
                        rtf_content = rtf_file.read()
                    break
                except (UnicodeDecodeError, LookupError):
                    continue

            # If all encodings fail, use errors='ignore' as fallback
            if rtf_content is None:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as rtf_file:
                    rtf_content = rtf_file.read()

            # Use striprtf to extract text
            text_content = rtf_to_text(rtf_content)
            result = text_content.strip() if text_content else ""
            if result:
                logger.info(
                    f"RTF extraction success: {os.path.basename(file_path)} ({len(result)} chars)"
                )
            return result
        except Exception as e:
            logger.error(f"Error extracting text from RTF {file_path}: {e}")
            return ""

    @staticmethod
    async def convert(file_path: str) -> str:
        """Convert RTF to text asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(_thread_pool, RTFConverter._extract_text, file_path)


class TextPassthrough(TextConverter):
    """Pass through text files (already in text format)."""

    @staticmethod
    def _read_text(file_path: str) -> str:
        """Read text file with encoding handling."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    result = f.read().strip()
                    if result:
                        logger.info(
                            f"Text read success: {os.path.basename(file_path)} ({len(result)} chars)"
                        )
                    return result
            except (UnicodeDecodeError, LookupError):
                continue

        # Fallback: ignore errors
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}")
            return ""

    @staticmethod
    async def convert(file_path: str) -> str:
        """Read text file asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(_thread_pool, TextPassthrough._read_text, file_path)


class FileConverter:
    """Main converter class that routes to appropriate converter based on file type."""

    _converters = {
        "pdf": PDFConverter,
        "word": WordConverter,
        "image": ImageConverter,
        "rtf": RTFConverter,
        "text": TextPassthrough,
    }

    @classmethod
    async def convert_to_text(cls, file_path: str) -> str:
        """
        Convert any supported file to text.

        Args:
            file_path: Path to the file to convert.

        Returns:
            Extracted text content.
        """
        extension = os.path.splitext(file_path)[1].lower()
        file_type = SupportedExtensions.get_file_type(extension)

        if file_type == "unknown":
            logger.warning(f"Unsupported file type: {extension}")
            return ""

        converter = cls._converters.get(file_type)
        if not converter:
            logger.error(f"No converter found for file type: {file_type}")
            return ""

        try:
            text = await converter.convert(file_path)
            return text
        except Exception as e:
            logger.error(f"Error converting file {file_path}: {e}")
            return ""

    @classmethod
    async def convert_batch(cls, file_paths: list[str], concurrency: int = 50) -> dict[str, str]:
        """
        Convert multiple files to text concurrently.

        Args:
            file_paths: List of file paths to convert.
            concurrency: Maximum number of concurrent conversions.

        Returns:
            Dictionary mapping file paths to extracted text.
        """
        semaphore = asyncio.Semaphore(concurrency)

        async def convert_with_semaphore(path: str) -> tuple[str, str]:
            async with semaphore:
                text = await cls.convert_to_text(path)
                return (path, text)

        tasks = [convert_with_semaphore(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        output = {}
        for result in results:
            if isinstance(result, BaseException):
                logger.error(f"Batch conversion error: {result}")
            elif isinstance(result, tuple):
                path, text = result
                output[path] = text

        return output
